#!/usr/bin/env python3
"""
Generate formal compliance memo as DOCX file.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from datetime import datetime

def create_compliance_memo():
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Header
    header = doc.add_heading('COMPLIANCE MEMORANDUM', 0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add memo metadata
    doc.add_paragraph()
    
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.style = 'Table Grid'
    
    meta_data = [
        ('TO:', 'Terminal Operations Review Board'),
        ('FROM:', 'Safety & Compliance Department'),
        ('DATE:', '29 March 2026'),
        ('SUBJECT:', 'IMO Circular Compliance Analysis and Automated Inspection System Implementation')
    ]
    
    for i, (label, value) in enumerate(meta_data):
        cell_label = meta_table.rows[i].cells[0]
        cell_value = meta_table.rows[i].cells[1]
        
        cell_label.text = label
        cell_value.text = value
        
        # Bold the labels
        for paragraph in cell_label.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Executive Summary
    doc.add_heading('1. EXECUTIVE SUMMARY', level=1)
    
    exec_summary = """This memorandum presents findings from a comprehensive cross-reference analysis of IMO safety circulars against vessel operational history for OPS-LISBON-VESSEL-908 (ATLANTIC PIONEER). The analysis covers the period from March 2024 through February 2026 and identifies critical compliance gaps requiring immediate remediation.

KEY FINDINGS:

• 20 safety incidents recorded over 19 voyages (63% voyage affect rate)
• 14 incidents (70%) directly addressable by new IMO circulars
• 5 compliance gaps require immediate action before next voyage
• €642,000 invested in circular-driven maintenance (65% of total spend)
• Estimated remediation cost for remaining gaps: €185,000 - €240,000

CRITICAL COMPLIANCE GAPS IDENTIFIED:

1. Gas monitor calibration expired (INC-908-2026-002) - OPEN status, 2 weeks overdue
2. IMSBC 08-25 voluntary application not followed (INC-908-2026-001)
3. Fumigation contractor procedures not updated per MSC.1/Circ.1264/Rev.1
4. Cargo Securing Manual amendment for lashing software v5.0 not verified
5. Dangerous goods shipper documentation gaps (3 incidents)

This memo recommends implementation of an automated Container Loading Flag System to prevent recurrence of documented incidents through real-time trigger-based inspection protocols."""
    
    for para in exec_summary.split('\n\n'):
        if para.strip():
            doc.add_paragraph(para)
    
    doc.add_paragraph()
    
    # Background
    doc.add_heading('2. BACKGROUND', level=1)
    
    background = """The International Maritime Organization (IMO) Maritime Safety Committee (MSC 110, June 2025) and Sub-Committee on Carriage of Cargoes and Containers (CCC 10, September 2024) issued significant safety circulars affecting cargo operations, dangerous goods handling, and vessel safety systems.

Vessel OPS-LISBON-VESSEL-908 (ATLANTIC PIONEER), an 8,500 TEU container ship operated by EuroContainer Lines, has been selected for detailed compliance analysis based on:

• High voyage frequency (19 voyages in 24 months)
• Diverse cargo portfolio (containers, bulk, dangerous goods, reefer)
• Multiple port state control regimes (Paris MoU, Tokyo MoU, Latin America MoU)
• Documented incident history across multiple safety categories

The vessel's operational history provides a representative sample for evaluating circular implementation effectiveness and identifying systemic compliance gaps."""
    
    for para in background.split('\n\n'):
        if para.strip():
            doc.add_paragraph(para)
    
    doc.add_paragraph()
    
    # Detailed Findings
    doc.add_heading('3. DETAILED FINDINGS', level=1)
    
    # 3.1 IMSBC Code
    doc.add_heading('3.1 IMSBC Code Amendments 08-25', level=2)
    
    doc.add_paragraph('Entry into Force: 1 January 2027 (voluntary from 1 January 2026)')
    doc.add_paragraph('Related Incidents: 4')
    
    findings_table = doc.add_table(rows=5, cols=4)
    findings_table.style = 'Table Grid'
    
    # Header row
    header_row = findings_table.rows[0]
    headers = ['Incident ID', 'Date', 'Issue', 'Status']
    for i, h in enumerate(headers):
        header_row.cells[i].text = h
        header_row.cells[i].paragraphs[0].runs[0].bold = True
    
    # Data rows
    incident_data = [
        ('INC-908-2024-007', '2024-11-03', 'DRI under old IMSBC schedule', 'Voluntary Apply'),
        ('INC-908-2025-005', '2025-06-28', 'Fish meal antioxidant undeclared', 'Voluntary Apply'),
        ('INC-908-2026-001', '2026-01-15', 'Zinc slag under generic schedule', 'MANDATORY'),
        ('INC-908-2024-008', '2024-12-18', 'Lashing software discrepancy', 'Voluntary Apply')
    ]
    
    for i, (inc_id, date, issue, status) in enumerate(incident_data, 1):
        row = findings_table.rows[i]
        row.cells[0].text = inc_id
        row.cells[1].text = date
        row.cells[2].text = issue
        row.cells[3].text = status
    
    doc.add_paragraph()
    doc.add_paragraph('COMPLIANCE GAP: Zinc slag incident (INC-908-2026-001) occurred AFTER voluntary application date of 1 January 2026. This constitutes a reportable deficiency if discovered during Port State Control inspection.', style='Intense Quote')
    
    doc.add_page_break()
    
    # 3.2 Fumigation Safety
    doc.add_heading('3.2 MSC.1/Circ.1264/Rev.1 - Fumigation Safety', level=2)
    
    doc.add_paragraph('Status: Approved MSC 110 (June 2025)')
    doc.add_paragraph('Related Incidents: 3')
    
    fum_table = doc.add_table(rows=4, cols=4)
    fum_table.style = 'Table Grid'
    
    header_row = fum_table.rows[0]
    headers = ['Incident ID', 'Date', 'Issue', 'Status']
    for i, h in enumerate(headers):
        header_row.cells[i].text = h
        header_row.cells[i].paragraphs[0].runs[0].bold = True
    
    fum_data = [
        ('INC-908-2024-002', '2024-04-22', 'Phosphine gas detection (12 ppm)', 'Procedural Gap'),
        ('INC-908-2025-003', '2025-03-30', 'Fumigation warning sign missing', 'Procedural Gap'),
        ('INC-908-2024-007', '2024-11-03', 'DRI cargo - phosphine fire risk', 'Procedural Gap')
    ]
    
    for i, (inc_id, date, issue, status) in enumerate(fum_data, 1):
        row = fum_table.rows[i]
        row.cells[0].text = inc_id
        row.cells[1].text = date
        row.cells[2].text = issue
        row.cells[3].text = status
    
    doc.add_paragraph()
    doc.add_paragraph('COMPLIANCE GAP: Fumigation contractor vetting procedure not updated per MSC.1/Circ.1264/Rev.1. Next fumigation operation carries elevated risk.', style='Intense Quote')
    
    # 3.3 Enclosed Space Entry
    doc.add_heading('3.3 Revised A.1050(27) - Enclosed Space Entry', level=2)
    
    doc.add_paragraph('Status: ADOPTED MSC 110 (June 2025)')
    doc.add_paragraph('Related Incidents: 2')
    
    es_table = doc.add_table(rows=3, cols=4)
    es_table.style = 'Table Grid'
    
    header_row = es_table.rows[0]
    headers = ['Incident ID', 'Date', 'Issue', 'Status']
    for i, h in enumerate(headers):
        header_row.cells[i].text = h
        header_row.cells[i].paragraphs[0].runs[0].bold = True
    
    es_data = [
        ('INC-908-2024-006', '2024-09-25', 'Oxygen deficiency (18.5%)', 'Partially Resolved'),
        ('INC-908-2026-002', '2026-02-08', 'Gas monitor calibration expired', 'OPEN GAP')
    ]
    
    for i, (inc_id, date, issue, status) in enumerate(es_data, 1):
        row = es_table.rows[i]
        row.cells[0].text = inc_id
        row.cells[1].text = date
        row.cells[2].text = issue
        row.cells[3].text = status
    
    doc.add_paragraph()
    doc.add_paragraph('CRITICAL COMPLIANCE GAP: Gas monitor calibration expired 2 weeks ago (INC-908-2026-002). Status still "Open". This is a detainable deficiency under Port State Control inspection. Revised A.1050(27) explicitly requires calibrated equipment for enclosed space entry.', style='Intense Quote')
    
    # 3.4 Cargo Securing
    doc.add_heading('3.4 MSC.1/Circ.1353/Rev.2 - Cargo Securing Manual', level=2)
    
    doc.add_paragraph('Status: Revision initiated CCC 10 (September 2024)')
    doc.add_paragraph('Related Incidents: 5 (highest frequency category - 26% incident rate)')
    
    doc.add_paragraph()
    doc.add_paragraph('COMPLIANCE GAP: Lashing software discrepancy (INC-908-2024-008) was "corrected" but CSM formal amendment for software v5.0 not verified with classification society. This is a common Port State Control deficiency.', style='Intense Quote')
    
    doc.add_page_break()
    
    # 3.5 Dangerous Goods
    doc.add_heading('3.5 IMDG Code Column 17 Amendments', level=2)
    
    doc.add_paragraph('Expected Adoption: MSC 111 (2026)')
    doc.add_paragraph('Related Incidents: 3')
    
    dg_table = doc.add_table(rows=4, cols=4)
    dg_table.style = 'Table Grid'
    
    header_row = dg_table.rows[0]
    headers = ['Incident ID', 'Date', 'Issue', 'Status']
    for i, h in enumerate(headers):
        header_row.cells[i].text = h
        header_row.cells[i].paragraphs[0].runs[0].bold = True
    
    dg_data = [
        ('INC-908-2024-003', '2024-05-08', 'Undeclared lithium batteries', 'Shipper Gap'),
        ('INC-908-2025-001', '2025-01-22', 'Column 17 information missing', 'Shipper Gap'),
        ('INC-908-2025-009', '2025-11-30', 'Lithium battery thermal runaway', 'Shipper Gap')
    ]
    
    for i, (inc_id, date, issue, status) in enumerate(dg_data, 1):
        row = dg_table.rows[i]
        row.cells[0].text = inc_id
        row.cells[1].text = date
        row.cells[2].text = issue
        row.cells[3].text = status
    
    doc.add_paragraph()
    
    # Compliance Summary
    doc.add_heading('4. COMPLIANCE GAP SUMMARY', level=1)
    
    summary_table = doc.add_table(rows=6, cols=5)
    summary_table.style = 'Table Grid'
    
    # Header
    header_row = summary_table.rows[0]
    headers = ['Priority', 'Gap', 'Circular', 'Incident', 'Remediation']
    for i, h in enumerate(headers):
        header_row.cells[i].text = h
        header_row.cells[i].paragraphs[0].runs[0].bold = True
    
    # Data
    summary_data = [
        ('CRITICAL', 'Gas monitor calibration expired', 'A.1050(27)', 'INC-908-2026-002', 'Verify recalibration; update tracking'),
        ('CRITICAL', 'IMSBC voluntary not applied', 'IMSBC 08-25', 'INC-908-2026-001', 'Update cargo acceptance checklist'),
        ('HIGH', 'Fumigation procedures outdated', 'MSC.1/Circ.1264/Rev.1', 'INC-908-2024-002', 'Update contractor vetting'),
        ('HIGH', 'CSM amendment not verified', 'MSC.1/Circ.1353/Rev.2', 'INC-908-2024-008', 'Confirm class approval'),
        ('MEDIUM', 'DG shipper documentation gaps', 'IMDG 43-26', 'Multiple', 'Create approved shipper list')
    ]
    
    for i, (priority, gap, circular, incident, remediation) in enumerate(summary_data, 1):
        row = summary_table.rows[i]
        row.cells[0].text = priority
        row.cells[1].text = gap
        row.cells[2].text = circular
        row.cells[3].text = incident
        row.cells[4].text = remediation
    
    doc.add_paragraph()
    
    # Proposed Automated Workflow
    doc.add_heading('5. PROPOSED AUTOMATED WORKFLOW: CONTAINER LOADING FLAG SYSTEM', level=1)
    
    workflow_intro = """To prevent recurrence of documented incidents and ensure real-time IMO circular compliance, we propose implementation of an automated Container Loading Flag System. This system integrates with existing Terminal Operating Systems (TOS) and activates manual inspection protocols based on predefined trigger conditions.

SYSTEM ARCHITECTURE:

The Flag System operates as a middleware layer between the Terminal Operating System and loading operations, evaluating each container against trigger conditions before loading approval is granted.

TRIGGER CATEGORIES:

The system monitors seven trigger categories corresponding to IMO circular requirements and historical incident patterns:"""
    
    for para in workflow_intro.split('\n\n'):
        if para.strip():
            doc.add_paragraph(para)
    
    # Trigger conditions table
    doc.add_paragraph()
    
    trigger_table = doc.add_table(rows=8, cols=3)
    trigger_table.style = 'Table Grid'
    
    # Header
    header_row = trigger_table.rows[0]
    headers = ['Category', 'Trigger Count', 'Priority Level']
    for i, h in enumerate(headers):
        header_row.cells[i].text = h
        header_row.cells[i].paragraphs[0].runs[0].bold = True
    
    # Data
    trigger_data = [
        ('Dangerous Goods (DG)', '4 triggers', 'CRITICAL'),
        ('Bulk Cargo (BC)', '5 triggers', 'HIGH'),
        ('Fumigation (FUM)', '4 triggers', 'CRITICAL'),
        ('Cargo Securing (CS)', '5 triggers', 'HIGH'),
        ('Enclosed Space (ES)', '4 triggers', 'CRITICAL'),
        ('Pilot Transfer (PT)', '3 triggers', 'HIGH'),
        ('Vessel History (VH)', '3 triggers', 'MEDIUM')
    ]
    
    for i, (category, count, priority) in enumerate(trigger_data, 1):
        row = trigger_table.rows[i]
        row.cells[0].text = category
        row.cells[1].text = count
        row.cells[2].text = priority
    
    doc.add_paragraph()
    doc.add_heading('5.1 EXACT TRIGGER CONDITIONS', level=2)
    
    doc.add_paragraph('The following trigger conditions activate manual inspection requirements. Each trigger is evaluated in real-time during container processing:')
    
    # Dangerous Goods triggers
    doc.add_heading('Dangerous Goods Triggers', level=3)
    
    dg_triggers = [
        ('DG-001', 'Missing IMDG Column 17 Data', 'shippingDocument.column17 == null OR column17.flashpoint == null OR column17.toxicityLevel == null', 'HOLD_LOADING', 'Chief Officer'),
        ('DG-002', 'Lithium Battery Shipment', 'cargo.unNumber IN [UN3480, UN3481, UN3090, UN3091] AND packagingGroup != VERIFIED', 'ENHANCED_INSPECTION', 'Chief Officer'),
        ('DG-003', 'Undeclared DG Detection', 'container.xrayAnomaly == true OR container.sniffTestPositive == true', 'QUARANTINE', 'Master'),
        ('DG-004', 'DG Shipper Not Approved', 'shipper.dgComplianceRating == null OR shipper.dgComplianceRating < 0.85', 'ENHANCED_DOCUMENT_CHECK', 'Chief Officer')
    ]
    
    dg_trigger_table = doc.add_table(rows=5, cols=5)
    dg_trigger_table.style = 'Table Grid'
    
    header_row = dg_trigger_table.rows[0]
    headers = ['ID', 'Trigger Name', 'Condition', 'Action', 'Approval']
    for i, h in enumerate(headers):
        header_row.cells[i].text = h
        header_row.cells[i].paragraphs[0].runs[0].bold = True
    
    for i, (tid, name, condition, action, approval) in enumerate(dg_triggers, 1):
        row = dg_trigger_table.rows[i]
        row.cells[0].text = tid
        row.cells[1].text = name
        row.cells[2].text = condition
        row.cells[3].text = action
        row.cells[4].text = approval
    
    doc.add_paragraph()
    
    # Bulk Cargo triggers
    doc.add_heading('Bulk Cargo Triggers', level=3)
    
    bc_triggers = [
        ('BC-001', 'IMSBC Schedule Mismatch', 'cargo.imsbcSchedule != CURRENT_08-25 OR == LEGACY', 'HOLD_LOADING', 'Chief Officer'),
        ('BC-002', 'DRI Moisture Certification', 'cargo.bcsnCode IN [DRI-A, DRI-B] AND moistureContent.certified == false', 'HOLD_LOADING', 'Chief Officer'),
        ('BC-003', 'Fish Meal Self-Heating', 'cargo.bcsnCode CONTAINS FISH MEAL AND antioxidantTreatment.declared == false', 'HOLD_LOADING', 'Master'),
        ('BC-004', 'New IMSBC Schedule Cargo', 'cargo.bcsnCode IN [Zinc Slag, Iron Ore Briquettes, etc.]', 'ENHANCED_INSPECTION', 'Chief Officer'),
        ('BC-005', 'Temperature Anomaly', 'cargo.temperature > 55°C OR temperatureRising == true', 'QUARANTINE', 'Master')
    ]
    
    bc_trigger_table = doc.add_table(rows=6, cols=5)
    bc_trigger_table.style = 'Table Grid'
    
    header_row = bc_trigger_table.rows[0]
    headers = ['ID', 'Trigger Name', 'Condition', 'Action', 'Approval']
    for i, h in enumerate(headers):
        header_row.cells[i].text = h
        header_row.cells[i].paragraphs[0].runs[0].bold = True
    
    for i, (tid, name, condition, action, approval) in enumerate(bc_triggers, 1):
        row = bc_trigger_table.rows[i]
        row.cells[0].text = tid
        row.cells[1].text = name
        row.cells[2].text = condition
        row.cells[3].text = action
        row.cells[4].text = approval
    
    doc.add_page_break()
    
    # Fumigation triggers
    doc.add_heading('Fumigation Triggers', level=3)
    
    fum_triggers = [
        ('FUM-001', 'Phosphine Without Seal Test', 'fumigantType == PHOSPHINE AND sealIntegrityTest != PASSED', 'HOLD_LOADING', 'Master'),
        ('FUM-002', 'Missing Warning Signs', 'container.fumigated == true AND warningSignsApplied == false', 'HOLD_LOADING', 'Chief Officer'),
        ('FUM-003', 'Gas Level Anomaly', 'phosphineLevel > 0.1 ppm OR oxygenLevel < 19.5%', 'EVACUATE_AND_VENTILATE', 'Master'),
        ('FUM-004', 'Contractor Not Approved', 'contractor.approved == false OR lastAudit < 2025-06-01', 'ENHANCED_VERIFICATION', 'Chief Officer')
    ]
    
    fum_trigger_table = doc.add_table(rows=5, cols=5)
    fum_trigger_table.style = 'Table Grid'
    
    header_row = fum_trigger_table.rows[0]
    headers = ['ID', 'Trigger Name', 'Condition', 'Action', 'Approval']
    for i, h in enumerate(headers):
        header_row.cells[i].text = h
        header_row.cells[i].paragraphs[0].runs[0].bold = True
    
    for i, (tid, name, condition, action, approval) in enumerate(fum_triggers, 1):
        row = fum_trigger_table.rows[i]
        row.cells[0].text = tid
        row.cells[1].text = name
        row.cells[2].text = condition
        row.cells[3].text = action
        row.cells[4].text = approval
    
    doc.add_paragraph()
    
    # Action definitions
    doc.add_heading('5.2 ACTION DEFINITIONS', level=2)
    
    action_table = doc.add_table(rows=6, cols=3)
    action_table.style = 'Table Grid'
    
    header_row = action_table.rows[0]
    headers = ['Action', 'Description', 'Escalation Time']
    for i, h in enumerate(headers):
        header_row.cells[i].text = h
        header_row.cells[i].paragraphs[0].runs[0].bold = True
    
    action_data = [
        ('HOLD_LOADING', 'Stop loading. Container cannot be loaded until inspection completed.', '30 minutes'),
        ('ENHANCED_INSPECTION', 'Enhanced inspection required. Loading may continue with inspection.', '1 hour'),
        ('QUARANTINE', 'Isolate container in quarantine area. Master approval required.', '15 minutes'),
        ('EVACUATE_AND_VENTILATE', 'Immediate evacuation. Activate emergency ventilation.', 'IMMEDIATE'),
        ('SUSPEND_OPERATIONS', 'Suspend all operations until equipment verified.', 'IMMEDIATE')
    ]
    
    for i, (action, desc, time) in enumerate(action_data, 1):
        row = action_table.rows[i]
        row.cells[0].text = action
        row.cells[1].text = desc
        row.cells[2].text = time
    
    doc.add_paragraph()
    
    # Implementation timeline
    doc.add_heading('6. IMPLEMENTATION TIMELINE', level=1)
    
    timeline = """PHASE 1 - IMMEDIATE (Within 48 hours):
• Verify gas monitor calibration certificates (INC-908-2026-002)
• Update cargo acceptance checklist for IMSBC 08-25
• Suspend enclosed space operations until calibration verified

PHASE 2 - SHORT TERM (Within 7 days):
• Update fumigation contractor procedures per MSC.1/Circ.1264/Rev.1
• Verify CSM amendment for lashing software v5.0 with classification society
• Deploy Flag System configuration to test environment

PHASE 3 - MEDIUM TERM (Within 30 days):
• Create DG shipper approved list with compliance ratings
• Implement calibration tracking system with automated alerts
• Complete Flag System integration with Terminal Operating System
• Conduct crew training on trigger response procedures

PHASE 4 - LONG TERM (Within 90 days):
• Full Flag System production deployment
• Monthly statistics reporting implementation
• Port State Control deficiency rate monitoring
• Quarterly system effectiveness review"""
    
    for para in timeline.split('\n\n'):
        if para.strip():
            doc.add_paragraph(para)
    
    doc.add_paragraph()
    
    # Recommendations
    doc.add_heading('7. RECOMMENDATIONS', level=1)
    
    recommendations = """1. APPROVE immediate implementation of the Container Loading Flag System as specified in the attached configuration file (container_loading_flag_system_config.json).

2. AUTHORIZE emergency recalibration of all gas monitoring equipment and suspension of enclosed space operations until certificates verified.

3. DIRECT Operations Department to update cargo acceptance procedures to require IMSBC 08-25 schedules for all bulk cargo, effective immediately.

4. MANDATE Safety Department to revise fumigation contractor vetting procedures per MSC.1/Circ.1264/Rev.1 before next fumigation operation.

5. REQUIRE Technical Superintendent to obtain classification society approval letter for CSM amendment incorporating lashing software v5.0 parameters.

6. ESTABLISH monthly Flag System statistics reporting to Terminal Operations Review Board, including trigger counts, resolution times, and repeat offender analysis.

7. ALLOCATE budget for Flag System integration with existing Terminal Operating System and gas monitoring infrastructure."""
    
    for para in recommendations.split('\n\n'):
        if para.strip():
            doc.add_paragraph(para)
    
    doc.add_paragraph()
    
    # Attachments
    doc.add_heading('8. ATTACHMENTS', level=1)
    
    attachments = """1. IMO Circular Cross-Reference Analysis (imo_circular_vessel_cross_reference_analysis.md)
2. Vessel Operational History (vessel_OPS-LISBON-VESSEL-908_operational_history.json)
3. Container Loading Flag System Configuration (container_loading_flag_system_config.json)
4. IMO Safety Circulars Summary (imo_safety_circulars_2026-03-29.md)"""
    
    for line in attachments.split('\n'):
        if line.strip():
            doc.add_paragraph(line, style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Signature block
    doc.add_paragraph('_' * 60)
    doc.add_paragraph()
    doc.add_paragraph('Prepared by: Safety & Compliance Department')
    doc.add_paragraph('Date: 29 March 2026')
    doc.add_paragraph()
    doc.add_paragraph('For questions regarding this memorandum, contact:')
    doc.add_paragraph('Safety Department | safety@eurocontainerlines.com')
    
    # Save document
    doc.save('/mnt/afs_toolcall/zhoulin3/.openclaw/workspaces/gendata-worker-27/compliance_memo_terminal_operations_review_board.docx')
    print("Document saved successfully!")

if __name__ == '__main__':
    create_compliance_memo()
